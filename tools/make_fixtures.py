#!/usr/bin/env python3
"""Generate Verbatim-compatible .docx test fixtures for Cardstock.

Every fixture embeds the REAL Verbatim 6.0.0 style definitions: the entire
<w:styles> part of each generated document is replaced, verbatim, with the
styles.xml extracted from Debate.dotm into spec/verbatim-styles.xml.
That file is GPL (from Verbatim) and is NOT committed; to regenerate fixtures,
download Verbatim's Debate.dotm, `unzip Debate.dotm word/styles.xml`, and place
it at spec/verbatim-styles.xml locally first.
Word/Verbatim therefore recognize Pocket/Hat/Block/Tag/Cite/Underline/
Emphasis exactly as the template defines them.

Style application is done by styleId (w:pStyle / w:rStyle) via docx.oxml so
that alias handling in python-docx's name lookup can never skew the output.

Usage: python3 tools/make_fixtures.py
"""

import copy
import random
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
STYLES_XML = ROOT / "spec" / "verbatim-styles.xml"
FIXTURES = ROOT / "tests" / "fixtures"

# ---------------------------------------------------------------- style layer

_VERBATIM_STYLES_ROOT = parse_xml(STYLES_XML.read_bytes())


def new_verbatim_document() -> Document:
    """A blank document whose styles part is byte-equivalent to Verbatim's."""
    doc = Document()
    root = doc.styles.element  # <w:styles> root of the styles part
    for child in list(root):
        root.remove(child)
    for attr, val in _VERBATIM_STYLES_ROOT.attrib.items():
        root.set(attr, val)
    for child in _VERBATIM_STYLES_ROOT:
        root.append(copy.deepcopy(child))
    # drop the stale python-docx default body paragraph? Document() starts
    # empty in python-docx >= 0.8, nothing to do.
    return doc


# Debate styleIds (see spec/verbatim-styles.md, section 1)
POCKET = "Heading1"       # w:name "heading 1", alias "Pocket"
HAT = "Heading2"          # w:name "heading 2", alias "Hat"
BLOCK = "Heading3"        # w:name "heading 3", alias "Block"
TAG = "Heading4"          # w:name "heading 4", alias "Tag"
CITE = "Style13ptBold"    # character, alias "Cite"
UNDERLINE = "StyleUnderline"  # character, alias "Underline"
EMPHASIS = "Emphasis"     # character, bold+underline (not italic)
HYPERLINK = "Hyperlink"


def para(doc, style_id=None):
    p = doc.add_paragraph()
    if style_id is not None:
        pPr = p._p.get_or_add_pPr()
        pStyle = pPr.get_or_add_pStyle()
        pStyle.val = style_id
    return p


def run(p, text, char_style=None, bold=None, underline=None, size_pt=None,
        highlight=None):
    r = p.add_run(text)
    if char_style is not None:
        rPr = r._r.get_or_add_rPr()
        rStyle = rPr.get_or_add_rStyle()
        rStyle.val = char_style
    if bold is not None:
        r.font.bold = bold
    if underline is not None:
        r.font.underline = underline
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    if highlight is not None:
        r.font.highlight_color = highlight
    return r


def add_hyperlink(p, url, text):
    """w:hyperlink with an external relationship + Hyperlink char style."""
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = p._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    r = hl.makeelement(qn("w:r"), {})
    rPr = r.makeelement(qn("w:rPr"), {})
    rStyle = rPr.makeelement(qn("w:rStyle"), {qn("w:val"): HYPERLINK})
    rPr.append(rStyle)
    r.append(rPr)
    t = r.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    p._p.append(hl)
    return hl


# --------------------------------------------------------------- card helpers

def add_cite(doc, author_date, detail):
    """Cite paragraph: bold 13pt author-date run + plain small detail run."""
    p = para(doc)  # Normal ("Normal/Card")
    run(p, author_date, char_style=CITE)
    run(p, " " + detail)
    return p


def add_card_body(doc, segments):
    """segments: list of (text, kind) where kind in
    {'plain', 'under', 'under_hl:<color>', 'emph', 'min'}."""
    p = para(doc)
    for text, kind in segments:
        if kind == "plain":
            run(p, text)
        elif kind == "under":
            run(p, text, char_style=UNDERLINE)
        elif kind.startswith("under_hl:"):
            color = getattr(WD_COLOR_INDEX, kind.split(":", 1)[1])
            run(p, text, char_style=UNDERLINE, highlight=color)
        elif kind == "emph":
            run(p, text, char_style=EMPHASIS)
        elif kind == "min":
            # Verbatim shrink convention: direct 8pt, no underline
            run(p, text, size_pt=8)
        else:
            raise ValueError(kind)
    return p


# ------------------------------------------------------------------- fixtures

def fx_01_minimal():
    doc = new_verbatim_document()
    run(para(doc, POCKET), "AT: Economy DA")
    run(para(doc, HAT), "Uniqueness Answers")
    run(para(doc, BLOCK), "AT: Economy Resilient")
    run(para(doc, TAG), "The economy is fragile — recent data proves")
    add_cite(doc, "Smith '24",
             "(John Smith, Professor of Economics at Example University, "
             "\"The Fragile Recovery,\" Journal of Things, 3-14-2024, "
             "https://example.com/fragile)")
    add_card_body(doc, [
        ("The consensus view holds that ", "plain"),
        ("the recovery rests on unusually weak foundations", "under"),
        (". Analysts who examined quarterly filings found reasons for "
         "concern that were widely underreported at the time, ", "min"),
        ("and every leading indicator now points the same direction",
         "under"),
        (".", "plain"),
    ])
    return doc


def fx_02_highlights():
    doc = new_verbatim_document()
    run(para(doc, TAG), "Highlight color coverage card")
    add_cite(doc, "Jones '25", "(citation detail for highlight fixture)")
    add_card_body(doc, [
        ("Lead-in text ", "plain"),
        ("yellow highlighted phrase", "under_hl:YELLOW"),
        (" middle underlined ", "under"),
        ("cyan highlighted phrase", "under_hl:TURQUOISE"),
        (" more underline ", "under"),
        ("green highlighted phrase", "under_hl:BRIGHT_GREEN"),
        (" and ", "under"),
        ("magenta highlighted phrase", "under_hl:PINK"),
        (" trailing plain text.", "plain"),
    ])
    return doc


def fx_03_emphasis():
    doc = new_verbatim_document()
    run(para(doc, TAG), "Emphasis and bold interaction card")
    add_cite(doc, "Nguyen '23", "(citation detail for emphasis fixture)")
    p = para(doc)
    run(p, "Plain lead-in ", )
    run(p, "underlined context ", char_style=UNDERLINE)
    run(p, "emphasized key words", char_style=EMPHASIS)
    run(p, " more underline ", char_style=UNDERLINE)
    run(p, "directly-bolded words", bold=True)
    run(p, " and ", char_style=UNDERLINE)
    run(p, "bold AND underlined direct formatting", bold=True, underline=True)
    run(p, " tail.", )
    return doc


def fx_04_structure():
    doc = new_verbatim_document()

    def card(tag_text, author):
        run(para(doc, TAG), tag_text)
        add_cite(doc, author, "(structure fixture citation)")
        add_card_body(doc, [
            ("Card body with ", "plain"),
            ("an underlined stretch of warranted text", "under"),
            (" and minimized filler around it.", "min"),
        ])

    # Pocket 1 with full hierarchy
    run(para(doc, POCKET), "Case Negative")
    run(para(doc, HAT), "Solvency Answers")
    run(para(doc, BLOCK), "AT: Federal Enforcement Solves")
    card("Enforcement fails — no resources", "Alvarez '22")
    card("Courts block implementation", "Baker '23")
    # loose plain paragraph between cards
    run(para(doc), "Loose analytic note between cards — plain Normal "
                   "paragraph, no heading style.")
    run(para(doc, BLOCK), "AT: States Fill In")
    card("Uniform action impossible", "Chen '24")

    # Hat with no pocket-level parent change, second hat
    run(para(doc, HAT), "Advantage Answers")
    card("Impact empirically denied", "Davis '21")  # tag directly under hat

    # Pocket 2
    run(para(doc, POCKET), "Topicality")
    # skipping levels: block with no hat
    run(para(doc, BLOCK), "T — Substantial (block with no hat above it)")
    card("Substantial means 90 percent", "Editor '20")

    # skipping levels the other way: hat with no pocket beneath a hatless flow
    run(para(doc, HAT), "Orphan Hat (document also opens levels out of order)")
    run(para(doc), "Trailing loose paragraph at end of document.")
    return doc


def fx_05_unicode():
    doc = new_verbatim_document()
    run(para(doc, POCKET), "Politics — «Übermensch» Pocket")
    run(para(doc, HAT), "“Smart Quotes” Hat — with em-dash")
    run(para(doc, BLOCK), "Café résumé block: ñ, é, ü")
    run(para(doc, TAG), "Tag with “curly quotes”, an em-dash — and ellipsis…")
    p = para(doc)
    run(p, "García-Márquez ’24 😎", char_style=CITE)
    run(p, " (María García-Márquez, señora of citations — “the definitive "
           "treatment”, piñata press, ", )
    add_hyperlink(p, "https://example.com/ünïcode?q=«test»",
                  "https://example.com/ünïcode")
    run(p, ")")
    add_card_body(doc, [
        ("Body text with non-ASCII: ", "plain"),
        ("“naïve” assumptions — étude of the débâcle", "under"),
        (" … and minimized text with emoji 🤝 and CJK 你好 mixed in.", "min"),
    ])
    return doc


def fx_06_empty():
    return new_verbatim_document()


_WORDS = ("policy economy hegemony deterrence escalation solvency uniqueness "
          "link internal impact warrant evidence fiat topicality inherency "
          "advantage disadvantage counterplan kritik alternative framework "
          "reduction treaty sanctions alliance credibility signal resolve "
          "miscalculation brinkmanship stability multipolarity proliferation "
          "extinction mindset epistemology ontology reps discourse securitize "
          "growth collapse resilience transition degrowth innovation").split()


def _sentence(rng, n):
    ws = [rng.choice(_WORDS) for _ in range(n)]
    return (" ".join(ws)).capitalize() + ". "


def fx_07_large():
    rng = random.Random(0xCA2D)
    doc = new_verbatim_document()
    card_no = 0
    for pk in range(1, 11):  # 10 pockets
        run(para(doc, POCKET), f"Pocket {pk:02d} — {_sentence(rng, 3)}")
        for ht in range(1, 3):  # 2 hats each
            run(para(doc, HAT), f"Hat {pk:02d}.{ht} {_sentence(rng, 3)}")
            for bl in range(1, 3):  # 2 blocks each
                run(para(doc, BLOCK),
                    f"Block {pk:02d}.{ht}.{bl} {_sentence(rng, 4)}")
                for _ in range(5):  # 5 cards each => 10*2*2*5 = 200 cards
                    card_no += 1
                    run(para(doc, TAG),
                        f"Card {card_no:03d}: {_sentence(rng, 8)}")
                    add_cite(
                        doc,
                        f"Author{card_no} '2{card_no % 10}",
                        f"({_sentence(rng, 18)}accessed 8-21-2026, "
                        f"https://example.com/{card_no})")
                    segs = []
                    for _ in range(16):  # alternating min/underline segments
                        segs.append((_sentence(rng, 28), "min"))
                        segs.append((_sentence(rng, 12), "under"))
                        segs.append((_sentence(rng, 4),
                                     "under_hl:YELLOW"))
                    add_card_body(doc, segs)
    return doc


FIXTURE_BUILDERS = {
    "01-minimal.docx": fx_01_minimal,
    "02-highlights.docx": fx_02_highlights,
    "03-emphasis.docx": fx_03_emphasis,
    "04-structure.docx": fx_04_structure,
    "05-unicode.docx": fx_05_unicode,
    "06-empty.docx": fx_06_empty,
    "07-large.docx": fx_07_large,
}


def main():
    FIXTURES.mkdir(parents=True, exist_ok=True)
    for name, builder in FIXTURE_BUILDERS.items():
        doc = builder()
        out = FIXTURES / name
        doc.save(out)
        print(f"wrote {out} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
